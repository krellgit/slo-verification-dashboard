#!/usr/bin/env python3
"""
Convert markdown documentation to DOCX format
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import sys

def add_heading(doc, text, level=1):
    """Add a heading with proper formatting"""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.runs[0].font.size = Pt(18)
        heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    elif level == 2:
        heading.runs[0].font.size = Pt(14)
        heading.runs[0].font.color.rgb = RGBColor(31, 41, 55)
    return heading

def add_paragraph(doc, text, bold=False, italic=False, code=False):
    """Add a paragraph with formatting"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if code:
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        p.paragraph_format.left_indent = Inches(0.25)

    return p

def add_code_block(doc, code):
    """Add a code block"""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Light gray background simulation with border
    p.style = 'Intense Quote'
    return p

def parse_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to DOCX"""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i].rstrip()

        # Code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block
                add_code_block(doc, '\n'.join(code_lines))
                code_lines = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Headings
        if line.startswith('# '):
            add_heading(doc, line[2:], level=1)
        elif line.startswith('## '):
            add_heading(doc, line[3:], level=2)
        elif line.startswith('### '):
            add_heading(doc, line[4:], level=3)
        elif line.startswith('#### '):
            add_heading(doc, line[5:], level=4)

        # Horizontal rules
        elif line.startswith('---'):
            doc.add_paragraph()  # Add spacing

        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            add_paragraph(doc, '• ' + line[2:])
        elif re.match(r'^\d+\.', line):
            add_paragraph(doc, line)

        # Bold/Italic/Status
        elif '**' in line:
            # Handle inline formatting
            parts = re.split(r'\*\*([^*]+)\*\*', line)
            p = doc.add_paragraph()
            for j, part in enumerate(parts):
                run = p.add_run(part)
                run.font.size = Pt(11)
                if j % 2 == 1:  # Bold parts
                    run.bold = True

        # Tables (simple markdown tables)
        elif '|' in line and not line.startswith('|---'):
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:
                # Check if this is first row of table
                if i + 1 < len(lines) and lines[i + 1].startswith('|---'):
                    # Create table
                    num_cols = len(cells)
                    table = doc.add_table(rows=1, cols=num_cols)
                    table.style = 'Light Grid Accent 1'

                    # Header row
                    hdr_cells = table.rows[0].cells
                    for j, cell_text in enumerate(cells):
                        hdr_cells[j].text = cell_text
                        for paragraph in hdr_cells[j].paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True

                    # Skip separator line
                    i += 2

                    # Add data rows
                    while i < len(lines) and '|' in lines[i]:
                        cells = [cell.strip() for cell in lines[i].split('|')[1:-1]]
                        if len(cells) == num_cols:
                            row_cells = table.add_row().cells
                            for j, cell_text in enumerate(cells):
                                row_cells[j].text = cell_text
                        i += 1
                    continue

        # Regular paragraphs
        elif line.strip():
            add_paragraph(doc, line)

        i += 1

    doc.save(docx_file)
    print(f"✅ Converted {md_file} → {docx_file}")

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python md-to-docx.py <input.md> [output.docx]")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else input_file.replace('.md', '.docx')

    parse_markdown_to_docx(input_file, output_file)
